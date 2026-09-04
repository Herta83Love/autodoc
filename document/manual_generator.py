# ============================================================================
# File: manual_generator.py
# ============================================================================

import json
import yaml
import ast
import re
import copy
import os
import shutil
import subprocess
import tempfile
import unicodedata
from datetime import datetime
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
from lxml import etree
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from document.ai_generator import (
    generate_manual_section
)

bookmark_id = 1
DOCUMENT_CONFIG = (
    "config/document.yaml"
)

ASSET_DIR = Path(__file__).resolve().parent / "assets"
LOGO_PATH = ASSET_DIR / "urmazi_logo.png"
COVER_PAGE_PATH = ASSET_DIR / "sentry_cover_page.png"

BRAND_BLUE = RGBColor(31, 101, 133)
BRAND_GREEN = RGBColor(112, 173, 71)
BRAND_ORANGE = RGBColor(247, 170, 20)
PALE_BLUE = RGBColor(227, 238, 242)
TEXT_BLACK = RGBColor(0, 0, 0)
MUTED_GRAY = RGBColor(100, 100, 100)
LIGHT_GRAY = "D9D9D9"


def convert_picture_to_page_anchor(inline, x, y):
    """Keep the cover artwork at a fixed position on the cover page."""

    inline.tag = qn("wp:anchor")

    for name, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "0",
        "behindDoc": "0",
        "locked": "1",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }.items():
        inline.set(name, value)

    extent = inline.find(qn("wp:extent"))
    simple_position = OxmlElement("wp:simplePos")
    simple_position.set("x", "0")
    simple_position.set("y", "0")

    horizontal = OxmlElement("wp:positionH")
    horizontal.set("relativeFrom", "page")
    horizontal_offset = OxmlElement("wp:posOffset")
    horizontal_offset.text = str(x)
    horizontal.append(horizontal_offset)

    vertical = OxmlElement("wp:positionV")
    vertical.set("relativeFrom", "page")
    vertical_offset = OxmlElement("wp:posOffset")
    vertical_offset.text = str(y)
    vertical.append(vertical_offset)

    wrap = OxmlElement("wp:wrapNone")
    insert_at = inline.index(extent)
    inline.insert(insert_at, simple_position)
    inline.insert(insert_at + 1, horizontal)
    inline.insert(insert_at + 2, vertical)
    inline.insert(insert_at + 4, wrap)

UI_TEXT = {
    "zh-TW": {
        "toc": "目錄 / TABLE OF CONTENTS",
        "document_info": "文件資訊",
        "product_name": "產品名稱",
        "version": "版本",
        "company": "公司",
        "generated_date": "產生日期",
        "revision_history": "修訂紀錄",
        "date": "日期",
        "description": "說明",
        "first_release": "文件首次產生",
        "introduction": "前言",
        "purpose": "文件目的",
        "audience": "適用對象",
        "login": "登入系統",
        "interface": "介面導覽",
        "action_icons": "畫面操作圖示",
        "icon": "圖示",
        "function": "功能說明",
        "overview": "功能概述",
        "business_value": "使用價值",
        "page_sections": "畫面組成",
        "fields": "欄位說明",
        "best_practices": "建議",
        "restrictions": "注意事項",
        "support_prefix": "文件意見與技術支援："
    },
    "en": {
        "toc": "TABLE OF CONTENTS",
        "document_info": "DOCUMENT INFORMATION",
        "product_name": "Product",
        "version": "Version",
        "company": "Company",
        "generated_date": "Generated Date",
        "revision_history": "REVISION HISTORY",
        "date": "Date",
        "description": "Description",
        "first_release": "Initial document generation",
        "introduction": "INTRODUCTION",
        "purpose": "Purpose",
        "audience": "Audience",
        "login": "Sign In",
        "interface": "Interface Overview",
        "action_icons": "SCREEN ACTION ICONS",
        "icon": "Icon",
        "function": "Function",
        "overview": "Overview",
        "business_value": "Business Value",
        "page_sections": "Screen Components",
        "fields": "Field Descriptions",
        "best_practices": "Recommendations",
        "restrictions": "Notes and Restrictions",
        "support_prefix": "Documentation and technical support: "
    }
}


def language_key(language):

    return "en" if str(language).lower().startswith("en") else "zh-TW"


def text_for(language, key):

    return UI_TEXT[language_key(language)][key]


def bilingual_name(local_name, english_name):

    local_name = str(local_name or "").strip()
    english_name = str(english_name or "").strip()

    if not english_name or local_name.casefold() == english_name.casefold():
        return local_name

    return f"{local_name}（{english_name}）"


def prepare_display_pages(pages, language):

    display_pages = copy.deepcopy(pages)

    if language_key(language) != "zh-TW":
        return display_pages

    for page in display_pages:
        page["category"] = bilingual_name(
            page.get("category"),
            page.get("english_category")
        )
        page["page"] = bilingual_name(
            page.get("page"),
            page.get("english_page")
        )
        page["tab"] = bilingual_name(
            page.get("tab"),
            page.get("english_tab")
        ) if page.get("tab") else None

    return display_pages


def set_run_font(
    run,
    name="Inter",
    size=None,
    color=None,
    bold=None,
    italic=None
):

    run.font.name = name

    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().get_or_add_rFonts()

    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    if size is not None:
        run.font.size = Pt(size)

    if color is not None:
        run.font.color.rgb = color

    if bold is not None:
        run.bold = bold

    if italic is not None:
        run.italic = italic


def add_field(run, instruction, cached_result="0"):
    """Insert a Word field using the standard multi-run OOXML structure.

    Keeping begin/instruction/separator/result/end in separate runs is
    important: Word and LibreOffice do not reliably refresh PAGEREF fields
    when all field elements are placed inside one run.
    """

    parent = run._r.getparent()
    run_index = parent.index(run._r)

    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run.append(begin)

    instruction_run = OxmlElement("w:r")
    field_code = OxmlElement("w:instrText")
    field_code.set(qn("xml:space"), "preserve")
    field_code.text = instruction
    instruction_run.append(field_code)

    separator_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separator_run.append(separate)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)

    run.text = str(cached_result)
    parent.insert(run_index, begin_run)
    parent.insert(run_index + 1, instruction_run)
    parent.insert(run_index + 2, separator_run)
    parent.insert(run_index + 4, end_run)


def add_paragraph_top_border(paragraph, color=LIGHT_GRAY, size="4"):

    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))

    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)

    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), size)
    top.set(qn("w:space"), "6")
    top.set(qn("w:color"), color)
    borders.append(top)


def configure_document_styles(document):

    background = document._element.find(qn("w:background"))

    if background is None:
        background = OxmlElement("w:background")
        document._element.insert(0, background)

    background.set(qn("w:color"), "FFFFFF")

    normal = document.styles["Normal"]
    normal.font.name = "Inter"
    normal.font.size = Pt(10)
    normal.font.color.rgb = TEXT_BLACK
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Inter")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Inter")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(4)

    heading_tokens = {
        "Heading 1": (18, TEXT_BLACK, 12, 4),
        "Heading 2": (14, TEXT_BLACK, 10, 2),
        "Heading 3": (11.5, TEXT_BLACK, 8, 2),
        "Heading 4": (10.5, TEXT_BLACK, 6, 2),
        "Heading 5": (10, TEXT_BLACK, 6, 2)
    }

    for name, (size, color, before, after) in heading_tokens.items():

        style = document.styles[name]
        style.font.name = "Inter"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Inter")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Inter")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name, left, size in [
        ("TOC 1", Cm(0), 10.5),
        ("TOC 2", Cm(0.55), 10),
        ("TOC 3", Cm(1.1), 9.5)
    ]:

        try:
            style = document.styles[name]
        except KeyError:
            continue

        style.font.name = "Inter"
        style.font.size = Pt(size)
        style.font.color.rgb = TEXT_BLACK
        style.paragraph_format.left_indent = left
        style.paragraph_format.space_after = Pt(2)


def enforce_white_background_black_text(document):
    """Write explicit white/black colors across every Word story and style."""

    background = document._element.find(qn("w:background"))

    if background is None:
        background = OxmlElement("w:background")
        document._element.insert(0, background)

    background.set(qn("w:color"), "FFFFFF")
    background.attrib.pop(qn("w:themeColor"), None)
    background.attrib.pop(qn("w:themeTint"), None)
    background.attrib.pop(qn("w:themeShade"), None)

    def force_run_black(run_element):
        r_pr = run_element.find(qn("w:rPr"))

        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            run_element.insert(0, r_pr)

        color = r_pr.find(qn("w:color"))

        if color is None:
            color = OxmlElement("w:color")
            r_pr.append(color)

        color.set(qn("w:val"), "000000")
        color.attrib.pop(qn("w:themeColor"), None)
        color.attrib.pop(qn("w:themeTint"), None)
        color.attrib.pop(qn("w:themeShade"), None)

    story_roots = [document._element]

    for section in document.sections:
        story_roots.extend([
            section.header._element,
            section.first_page_header._element,
            section.even_page_header._element,
            section.footer._element,
            section.first_page_footer._element,
            section.even_page_footer._element,
        ])

    seen_roots = set()

    for root in story_roots:
        root_id = id(root)

        if root_id in seen_roots:
            continue

        seen_roots.add(root_id)

        for run_element in root.iter(qn("w:r")):
            force_run_black(run_element)

        for shading in root.iter(qn("w:shd")):
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), "FFFFFF")
            shading.attrib.pop(qn("w:themeFill"), None)
            shading.attrib.pop(qn("w:themeFillTint"), None)
            shading.attrib.pop(qn("w:themeFillShade"), None)

        for cell_properties in root.iter(qn("w:tcPr")):
            if cell_properties.find(qn("w:shd")) is None:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), "FFFFFF")
                cell_properties.append(shading)

        for highlight in root.iter(qn("w:highlight")):
            highlight.set(qn("w:val"), "none")

    styles_root = document.styles._element

    for style in styles_root.iter(qn("w:style")):
        r_pr = style.find(qn("w:rPr"))

        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            style.append(r_pr)

        color = r_pr.find(qn("w:color"))

        if color is None:
            color = OxmlElement("w:color")
            r_pr.append(color)

        color.set(qn("w:val"), "000000")
        color.attrib.pop(qn("w:themeColor"), None)
        color.attrib.pop(qn("w:themeTint"), None)
        color.attrib.pop(qn("w:themeShade"), None)

    for color in styles_root.iter(qn("w:color")):
        color.set(qn("w:val"), "000000")
        color.attrib.pop(qn("w:themeColor"), None)
        color.attrib.pop(qn("w:themeTint"), None)
        color.attrib.pop(qn("w:themeShade"), None)

    for shading in styles_root.iter(qn("w:shd")):
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "FFFFFF")
        shading.attrib.pop(qn("w:themeFill"), None)
        shading.attrib.pop(qn("w:themeFillTint"), None)
        shading.attrib.pop(qn("w:themeFillShade"), None)

    for highlight in styles_root.iter(qn("w:highlight")):
        highlight.set(qn("w:val"), "none")

    doc_defaults = styles_root.find(qn("w:docDefaults"))

    if doc_defaults is not None:
        r_pr_default = doc_defaults.find(qn("w:rPrDefault"))

        if r_pr_default is None:
            r_pr_default = OxmlElement("w:rPrDefault")
            doc_defaults.append(r_pr_default)

        r_pr = r_pr_default.find(qn("w:rPr"))

        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            r_pr_default.append(r_pr)

        color = r_pr.find(qn("w:color"))

        if color is None:
            color = OxmlElement("w:color")
            r_pr.append(color)

        color.set(qn("w:val"), "000000")
        color.attrib.pop(qn("w:themeColor"), None)
        color.attrib.pop(qn("w:themeTint"), None)
        color.attrib.pop(qn("w:themeShade"), None)


def configure_page(section):

    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.55)


def clear_container(container):

    for paragraph in container.paragraphs:
        paragraph._element.getparent().remove(paragraph._element)

    for table in container.tables:
        table._element.getparent().remove(table._element)


def set_branded_header_footer(section):

    configure_page(section)

    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_container(section.header)
    clear_container(section.footer)

    header_p = section.header.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(4)

    if LOGO_PATH.exists():
        header_p.add_run().add_picture(
            str(LOGO_PATH),
            width=Inches(1.35)
        )

    footer_p = section.footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(4)
    add_paragraph_top_border(footer_p)

    # STYLEREF shows an error on front-matter pages before the first Heading 1
    # in Microsoft Word. Use a stable product label instead.
    section_run = footer_p.add_run("SENTRY PDNS")
    set_run_font(section_run, size=8.5, color=MUTED_GRAY)

    divider = footer_p.add_run("  |  ")
    set_run_font(divider, size=8.5, color=MUTED_GRAY)

    page_run = footer_p.add_run()
    add_field(page_run, " PAGE ", "1")
    set_run_font(page_run, size=8.5, color=MUTED_GRAY, bold=True)


def set_blank_header_footer(section):

    configure_page(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_container(section.header)
    clear_container(section.footer)
    section.header.add_paragraph()
    section.footer.add_paragraph()
def enable_field_updates(document):

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))

    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)

    update_fields.set(qn("w:val"), "true")


def find_command(name, macos_app_path=None):
    """Find a command in PATH or at a conventional macOS app location."""

    command = shutil.which(name)

    if command:
        return command

    if macos_app_path:
        app_command = Path(macos_app_path)

        if app_command.is_file():
            return str(app_command)

    return None


def normalize_pdf_line(text):
    return " ".join(
        unicodedata.normalize("NFKC", str(text or "")).split()
    )


def build_toc_targets(grouped):
    targets = []

    for category_index, (category, page_dict) in enumerate(grouped.items()):
        for page_index, (page_name, items) in enumerate(page_dict.items()):
            targets.append((
                toc_bookmark_name("p", category_index, page_index),
                page_name,
            ))
            rendered_tabs = set()

            for page in items:
                tab = (page.get("tab") or "").strip()

                if not tab or tab in rendered_tabs:
                    continue

                tab_index = len(rendered_tabs)
                targets.append((
                    toc_bookmark_name(
                        "t", category_index, page_index, tab_index
                    ),
                    tab,
                ))
                rendered_tabs.add(tab)

    return targets


def extract_pdf_page_texts(pdf_path, pdfinfo, pdftotext):
    info = subprocess.run(
        [pdfinfo, str(pdf_path)],
        text=True,
        capture_output=True,
        timeout=60,
        check=True,
    ).stdout
    match = re.search(r"^Pages:\s+(\d+)\s*$", info, re.MULTILINE)

    if not match:
        raise RuntimeError("無法取得 LibreOffice 輸出的 PDF 頁數")

    page_texts = []

    for page_number in range(1, int(match.group(1)) + 1):
        text = subprocess.run(
            [
                pdftotext,
                "-f", str(page_number),
                "-l", str(page_number),
                "-layout",
                str(pdf_path),
                "-",
            ],
            text=True,
            capture_output=True,
            timeout=60,
            check=True,
        ).stdout
        lines = {
            normalize_pdf_line(line)
            for line in text.splitlines()
            if normalize_pdf_line(line)
        }
        page_texts.append(lines)

    return page_texts


def match_toc_pages(grouped, page_texts):
    pages_by_bookmark = {}
    search_from = 0

    for bookmark_name, title in build_toc_targets(grouped):
        normalized_title = normalize_pdf_line(title)
        found_page = None

        for page_index in range(search_from, len(page_texts)):
            if normalized_title in page_texts[page_index]:
                found_page = page_index + 1
                break

        if found_page is None:
            raise RuntimeError(
                f"無法在排版結果中找到目錄標題：{title}"
            )

        pages_by_bookmark[bookmark_name] = found_page
        search_from = found_page - 1

    return pages_by_bookmark


def write_pageref_cached_values(docx_path, pages_by_bookmark):
    """Write resolved page numbers into the visible PAGEREF result runs."""

    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    }
    w_namespace = f"{{{namespaces['w']}}}"

    with ZipFile(docx_path, "r") as source:
        document_xml = source.read("word/document.xml")
        root = etree.fromstring(document_xml)

        updated = set()

        for paragraph in root.xpath(".//w:p", namespaces=namespaces):
            active_bookmark = None
            after_separator = False

            for run in paragraph.xpath("./w:r", namespaces=namespaces):
                instruction = "".join(
                    node.text or ""
                    for node in run.xpath("./w:instrText", namespaces=namespaces)
                )
                match = re.search(
                    r'PAGEREF\s+"?([^"\\\s]+)',
                    instruction,
                    re.IGNORECASE,
                )

                if match:
                    active_bookmark = match.group(1)

                for field_char in run.xpath(
                    "./w:fldChar", namespaces=namespaces
                ):
                    field_type = field_char.get(
                        f"{w_namespace}fldCharType"
                    )

                    if active_bookmark and field_type == "separate":
                        after_separator = True
                    elif active_bookmark and field_type == "end":
                        active_bookmark = None
                        after_separator = False

                if active_bookmark and after_separator:
                    text_nodes = run.xpath("./w:t", namespaces=namespaces)

                    if text_nodes:
                        page_number = pages_by_bookmark.get(active_bookmark)

                        if page_number is not None:
                            text_nodes[0].text = str(page_number)
                            updated.add(active_bookmark)

        missing = set(pages_by_bookmark) - updated

        if missing:
            raise RuntimeError(
                "無法寫入部分目錄頁碼：" + ", ".join(sorted(missing))
            )

        replacement_xml = etree.tostring(
            root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

        temp_fd, temp_name = tempfile.mkstemp(
            suffix=".docx",
            dir=str(Path(docx_path).resolve().parent),
        )
        os.close(temp_fd)

        try:
            with ZipFile(temp_name, "w", ZIP_DEFLATED) as target:
                for item in source.infolist():
                    data = (
                        replacement_xml
                        if item.filename == "word/document.xml"
                        else source.read(item.filename)
                    )
                    target.writestr(item, data)
        except Exception:
            Path(temp_name).unlink(missing_ok=True)
            raise

    os.replace(temp_name, docx_path)


def materialize_toc_page_numbers(output_path, grouped):
    """Resolve TOC pages with LibreOffice and store them in the DOCX."""

    soffice = find_command(
        "soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    )
    pdfinfo = find_command("pdfinfo")
    pdftotext = find_command("pdftotext")

    missing = [
        name
        for name, command in (
            ("LibreOffice", soffice),
            ("pdfinfo", pdfinfo),
            ("pdftotext", pdftotext),
        )
        if not command
    ]

    if missing:
        raise RuntimeError(
            "無法計算目錄頁碼，缺少："
            + ", ".join(missing)
            + "。請依照 readme.txt 安裝 LibreOffice 與 Poppler。"
        )

    print("🔄 正在使用 LibreOffice 計算目錄頁碼...")

    with tempfile.TemporaryDirectory(prefix="autodoc_toc_") as temp_dir:
        source_path = Path(output_path).resolve()
        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", temp_dir,
                str(source_path),
            ],
            text=True,
            capture_output=True,
            timeout=300,
            check=False,
        )

        pdf_path = Path(temp_dir) / f"{source_path.stem}.pdf"

        if result.returncode != 0 or not pdf_path.is_file():
            details = (result.stderr or result.stdout or "未知錯誤").strip()
            raise RuntimeError(
                "LibreOffice 無法完成文件分頁。\n" + details
            )

        page_texts = extract_pdf_page_texts(
            pdf_path,
            pdfinfo,
            pdftotext,
        )
        pages_by_bookmark = match_toc_pages(grouped, page_texts)

    write_pageref_cached_values(output_path, pages_by_bookmark)
    print(f"✅ 已寫入 {len(pages_by_bookmark)} 筆目錄頁碼")


def toc_bookmark_name(level, category_index, page_index=None, tab_index=None):
    """Return a short, unique, ASCII-only Word bookmark name."""

    parts = ["toc", level, str(category_index)]

    if page_index is not None:
        parts.append(str(page_index))

    if tab_index is not None:
        parts.append(str(tab_index))

    return "_".join(parts)


def add_toc_page_reference(paragraph, bookmark_name):
    """Add a right-aligned dotted leader and a clickable page reference."""

    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.clear_all()
    tab_stops.add_tab_stop(
        Cm(16.0),
        WD_TAB_ALIGNMENT.RIGHT,
        WD_TAB_LEADER.DOTS,
    )

    paragraph.add_run("\t")
    page_run = paragraph.add_run()
    add_field(page_run, f' PAGEREF "{bookmark_name}" \\h ')
    set_run_font(page_run, size=9.5, color=TEXT_BLACK)


def generate_manual_toc(document, grouped, language="zh-TW"):

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(24)
    title_run = title.add_run(text_for(language, "toc"))
    set_run_font(
        title_run,
        size=18,
        color=TEXT_BLACK,
        bold=True
    )

    for category_index, (category, page_dict) in enumerate(grouped.items()):

        category_bookmark = toc_bookmark_name("c", category_index)

        category_p = document.add_paragraph()
        category_p.paragraph_format.space_before = Pt(6)
        category_p.paragraph_format.space_after = Pt(2)
        add_internal_link(
            category_p,
            category,
            category_bookmark,
            color="000000",
            underline=False,
            bold=True,
            size=11
        )

        for page_index, (page_name, items) in enumerate(page_dict.items()):

            page_bookmark = toc_bookmark_name(
                "p", category_index, page_index
            )

            page_p = document.add_paragraph()
            page_p.paragraph_format.left_indent = Cm(0.55)
            page_p.paragraph_format.space_after = Pt(1)
            add_internal_link(
                page_p,
                page_name,
                page_bookmark,
                color="000000",
                underline=False,
                size=10
            )
            add_toc_page_reference(
                page_p,
                page_bookmark
            )

            rendered_tabs = set()

            for page in items:

                tab = (page.get("tab") or "").strip()

                if not tab or tab in rendered_tabs:
                    continue

                tab_index = len(rendered_tabs)
                tab_bookmark = toc_bookmark_name(
                    "t", category_index, page_index, tab_index
                )

                tab_p = document.add_paragraph()
                tab_p.paragraph_format.left_indent = Cm(1.15)
                tab_p.paragraph_format.space_after = Pt(1)
                add_internal_link(
                    tab_p,
                    tab,
                    tab_bookmark,
                    color="404040",
                    underline=False,
                    size=9.5
                )
                add_toc_page_reference(
                    tab_p,
                    tab_bookmark
                )
                rendered_tabs.add(tab)

    document.add_page_break()
def normalize_ai_content(data):

    if not data:
        return data

    if isinstance(data, str):

        text = data.strip()

        if (
            text.startswith("{")
            and text.endswith("}")
        ):
            try:

                obj = ast.literal_eval(text)

                rows = []

                for k, v in obj.items():

                    rows.append(
                        f"{k}\n{v}"
                    )

                return rows

            except Exception:
                return text

        return text

    if isinstance(data, list):

        result = []

        for item in data:

            if not isinstance(item, str):
                result.append(item)
                continue

            text = item.strip()

            if (
                text.startswith("{")
                and text.endswith("}")
            ):
                try:

                    obj = ast.literal_eval(text)

                    for k, v in obj.items():

                        result.append(
                            f"{k}\n{v}"
                        )

                except Exception:
                    result.append(text)

            else:
                result.append(text)

        return result

    return data

def load_document_config():

    with open(
        DOCUMENT_CONFIG,
        "r",
        encoding="utf-8"
    ) as f:

        return yaml.safe_load(
            f
        )


def get_localized_config(config, section, language):

    localized = (
        config.get("locales", {})
        .get(language_key(language), {})
        .get(section)
    )

    return localized or config[section]



def add_bookmark(
    paragraph,
    name
):

    global bookmark_id

    start = OxmlElement(
        "w:bookmarkStart"
    )

    start.set(
        qn("w:id"),
        str(bookmark_id)
    )

    start.set(
        qn("w:name"),
        name
    )

    end = OxmlElement(
        "w:bookmarkEnd"
    )

    end.set(
        qn("w:id"),
        str(bookmark_id)
    )

    # w:pPr must remain the first child of w:p. Placing a bookmark before it
    # creates invalid OOXML that LibreOffice tolerates but Word may ignore.
    p_pr = paragraph._p.pPr
    bookmark_position = 1 if p_pr is not None else 0
    paragraph._p.insert(
        bookmark_position,
        start
    )

    paragraph._p.append(
        end
    )

    bookmark_id += 1

def add_bullet_list(
    document,
    items
):

    for item in items:

        if not item:
            continue

        if isinstance(item, dict):

            if (
                "field" in item
                and "description" in item
            ):
                text = (
                    f"{item['field']}："
                    f"{item['description']}"
                )

            elif (
                "name" in item
                and "description" in item
            ):
                text = (
                    f"{item['name']}："
                    f"{item['description']}"
                )

            else:
                text = json.dumps(
                    item,
                    ensure_ascii=False
                )

        else:

            text = str(item)

        document.add_paragraph(
            text,
            style="List Bullet"
        )


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))

    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for side, value in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end)
    ]:
        node = tc_mar.find(qn(f"w:{side}"))

        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)

        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_table_row(row, repeat_header=False):

    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)

    if repeat_header:
        table_header = OxmlElement("w:tblHeader")
        table_header.set(qn("w:val"), "true")
        tr_pr.append(table_header)


def add_action_section(
    document,
    actions,
    button_descriptions,
    language="zh-TW"
):

    descriptions = {}

    for item in button_descriptions or []:

        if not isinstance(item, dict):
            continue

        description = str(item.get("description") or "").strip()
        index = item.get("button_index")

        if description and isinstance(index, int):
            descriptions[index] = description

    renderable_actions = []

    for index, action in enumerate(actions):

        image = action.get("image")
        description = descriptions.get(index)

        if not image or not description or not Path(image).is_file():
            continue

        renderable_actions.append((image, description))

    if not renderable_actions:
        return

    document.add_heading(
        text_for(language, "action_icons"),
        level=5
    )

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(2.6)
    table.columns[1].width = Cm(13.2)

    header = table.rows[0].cells
    configure_table_row(table.rows[0], repeat_header=True)
    header[0].text = text_for(language, "icon")
    header[1].text = text_for(language, "function")

    for cell, width in zip(header, [Cm(2.6), Cm(13.2)]):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)

        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9.5, color=TEXT_BLACK, bold=True)

    for image, description in renderable_actions:

        row = table.add_row()
        configure_table_row(row)
        cells = row.cells
        cells[0].width = Cm(2.6)
        cells[1].width = Cm(13.2)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cells[0])
        set_cell_margins(cells[1])

        image_p = cells[0].paragraphs[0]
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_p.paragraph_format.space_before = Pt(3)
        image_p.paragraph_format.space_after = Pt(3)

        try:
            image_p.add_run().add_picture(
                image,
                width=Inches(0.68)
            )
        except Exception as exc:
            print(f"按鈕圖片加入失敗: {image}: {exc}")
            continue

        description_p = cells[1].paragraphs[0]
        description_p.paragraph_format.space_before = Pt(3)
        description_p.paragraph_format.space_after = Pt(3)
        description_run = description_p.add_run(description)
        set_run_font(description_run, size=9.5, color=TEXT_BLACK)

    after_table = document.add_paragraph()
    after_table.paragraph_format.space_after = Pt(2)

def add_internal_link(
    paragraph,
    text,
    bookmark_name,
    color="0563C1",
    underline=True,
    bold=False,
    size=None
):

    hyperlink = OxmlElement(
        "w:hyperlink"
    )

    hyperlink.set(
        qn("w:anchor"),
        bookmark_name
    )

    run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")

    color_element = OxmlElement(
        "w:color"
    )

    color_element.set(
        qn("w:val"),
        color
    )

    underline_element = OxmlElement(
        "w:u"
    )

    underline_element.set(
        qn("w:val"),
        "single" if underline else "none"
    )

    rPr.append(color_element)
    rPr.append(underline_element)

    if bold:
        bold_element = OxmlElement("w:b")
        rPr.append(bold_element)

    if size is not None:
        size_element = OxmlElement("w:sz")
        size_element.set(qn("w:val"), str(int(size * 2)))
        rPr.append(size_element)

    run.append(rPr)

    text_elem = OxmlElement(
        "w:t"
    )

    text_elem.text = text

    run.append(text_elem)

    hyperlink.append(run)

    paragraph._p.append(
        hyperlink
    )

def group_pages(pages):

    result = {}

    for page in pages:

        category = page.get(
            "category",
            "Others"
        )

        page_name = page.get(
            "page",
            ""
        )

        result.setdefault(
            category,
            {}
        )

        result[
            category
        ].setdefault(
            page_name,
            []
        )

        result[
            category
        ][
            page_name
        ].append(
            page
        )

    return result


def add_cover(
    document,
    config,
    language="zh-TW"
):
    """Insert the approved PDF cover as one full-page image."""

    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)
    set_blank_header_footer(section)

    if not COVER_PAGE_PATH.is_file():
        raise FileNotFoundError(
            f"找不到封面圖片：{COVER_PAGE_PATH}"
        )

    cover_p = document.add_paragraph()
    cover_p.paragraph_format.space_before = Pt(0)
    cover_p.paragraph_format.space_after = Pt(0)
    cover_p.paragraph_format.line_spacing = 1
    picture = cover_p.add_run().add_picture(
        str(COVER_PAGE_PATH),
        width=Cm(21)
    )
    convert_picture_to_page_anchor(
        picture._inline,
        x=0,
        y=0
    )


def add_document_info(
    document,
    config,
    language="zh-TW"
):

    document.add_heading(text_for(language, "document_info"), level=1)

    document.add_paragraph(
        f"{text_for(language, 'product_name')}: {config['document']['product_name']}"
    )

    document.add_paragraph(
        f"{text_for(language, 'version')}: {config['document']['version']}"
    )

    document.add_paragraph(
        f"{text_for(language, 'company')}: {config['document']['company']}"
    )

    document.add_paragraph(
        f"{text_for(language, 'generated_date')}: {datetime.now().strftime('%Y-%m-%d')}"
    )

    publication = get_localized_config(config, "publication", language)

    notice = document.add_paragraph()
    notice.paragraph_format.space_before = Pt(18)
    notice_run = notice.add_run(publication["notice_title"])
    set_run_font(notice_run, size=14, color=TEXT_BLACK, bold=True)

    document.add_paragraph(
        publication["change_notice"]
    )

    document.add_paragraph(
        publication["trademark_notice"]
    )

    support = document.add_paragraph()
    support.add_run(text_for(language, "support_prefix")).bold = True
    support.add_run(publication["support_url"])

def add_revision_history(
    document,
    language="zh-TW"
):

    document.add_heading(
        text_for(language, "revision_history"),
        level=1
    )

    table = document.add_table(
        rows=2,
        cols=3
    )

    table.style = (
        "Table Grid"
    )

    table.autofit = False

    for cell, width in zip(table.rows[0].cells, [Cm(2.5), Cm(3.5), Cm(10)]):
        cell.width = width

    table.cell(
        0,
        0
    ).text = text_for(language, "version")

    table.cell(
        0,
        1
    ).text = text_for(language, "date")

    table.cell(
        0,
        2
    ).text = text_for(language, "description")

    table.cell(
        1,
        0
    ).text = "1.0"

    table.cell(
        1,
        1
    ).text = datetime.now().strftime(
        "%Y-%m-%d"
    )

    table.cell(
        1,
        2
    ).text = text_for(language, "first_release")

    for row in table.rows:
        for cell, width in zip(row.cells, [Cm(2.5), Cm(3.5), Cm(10)]):
            cell.width = width

    for cell in table.rows[0].cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "FFFFFF")
        cell._tc.get_or_add_tcPr().append(shading)

        for run in cell.paragraphs[0].runs:
            set_run_font(
                run,
                size=9.5,
                color=TEXT_BLACK,
                bold=True
            )

    for cell in table.rows[1].cells:
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9.5, color=TEXT_BLACK)

    document.add_page_break()

def sanitize_bookmark_name(text):

    return (
        str(text)
        .replace(" ", "_")
        .replace("&", "_")
        .replace("/", "_")
        .replace("\\", "_")
        .replace(":", "_")
        .replace("(", "_")
        .replace(")", "_")
    )
def add_introduction(
    document,
    config,
    language="zh-TW"
):

    intro = get_localized_config(config, "introduction", language)

    document.add_heading(intro["heading"], level=1)

    for section_key in ["what_is_sentry", "target_version", "logging_in"]:
        section_data = intro[section_key]
        heading = document.add_heading(section_data["title"], level=2)

        if section_key == "logging_in":
            heading.paragraph_format.page_break_before = True

        for paragraph_text in section_data["paragraphs"]:
            paragraph = document.add_paragraph(paragraph_text)

            if paragraph_text == config["document"].get("build_version"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.runs[0].bold = True

    document.add_page_break()


def add_back_cover(document, config, language="zh-TW"):

    back_cover = get_localized_config(config, "back_cover", language)
    publication = get_localized_config(config, "publication", language)

    logo_p = document.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_p.paragraph_format.space_after = Pt(110)

    if LOGO_PATH.exists():
        logo_p.add_run().add_picture(
            str(LOGO_PATH),
            width=Inches(1.65)
        )

    brand_p = document.add_paragraph()
    brand_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_p.paragraph_format.space_after = Pt(10)
    brand_run = brand_p.add_run("SENTRY PDNS")
    set_run_font(
        brand_run,
        size=28,
        color=BRAND_BLUE,
        bold=True
    )

    thanks_p = document.add_paragraph()
    thanks_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    thanks_p.paragraph_format.space_after = Pt(36)
    thanks_run = thanks_p.add_run(back_cover["thank_you"])
    set_run_font(thanks_run, size=14, color=TEXT_BLACK)

    support_p = document.add_paragraph()
    support_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    support_run = support_p.add_run(back_cover["support_label"])
    set_run_font(support_run, size=10, color=MUTED_GRAY, bold=True)

    link_p = document.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link_p.paragraph_format.space_after = Pt(110)
    link_run = link_p.add_run(publication["support_url"])
    set_run_font(link_run, size=10.5, color=BRAND_BLUE)

    sentry_p = document.add_paragraph()
    sentry_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentry_p.paragraph_format.space_after = Pt(16)
    sentry_run = sentry_p.add_run("SENTRY")
    set_run_font(sentry_run, size=46, color=PALE_BLUE, bold=True)

    copyright_p = document.add_paragraph()
    copyright_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_run = copyright_p.add_run(
        config["footer"]["copyright"]
    )
    set_run_font(copyright_run, size=8.5, color=MUTED_GRAY)


def generate_docx(
    pages=None,
    language="zh-TW",
    output_path=None
):

    global bookmark_id
    bookmark_id = 1

    config = load_document_config()

    if pages is None:
        metadata_path = (
            f"output/metadata_{language}.json"
            if Path(f"output/metadata_{language}.json").is_file()
            else "output/metadata.json"
        )

        with open(metadata_path, "r", encoding="utf-8") as f:
            pages = json.load(f)

    output_path = output_path or f"output/SENTRY_Manual_{language}.docx"
    pages = prepare_display_pages(pages, language)

    grouped = group_pages(
        pages
    )

    document = Document()

    enable_field_updates(document)
    configure_document_styles(document)

    add_cover(
        document,
        config,
        language
    )

    body_section = document.add_section(
        WD_SECTION.NEW_PAGE
    )
    set_branded_header_footer(body_section)

    add_document_info(
        document,
        config,
        language
    )

    add_revision_history(
        document,
        language
    )
    generate_manual_toc(document, grouped, language)

    add_introduction(
        document,
        config,
        language
    )

    for category_index, (category, page_dict) in enumerate(grouped.items()):

        h = document.add_heading(
            category,
            level=1
        )

        add_bookmark(
            h,
            toc_bookmark_name("c", category_index)
        )

        rendered_pages = set()

        for page_index, (page_name, items) in enumerate(page_dict.items()):

            if (
                page_name
                and page_name not in rendered_pages
            ):
                h = document.add_heading(
                    page_name,
                    level=2
                )

                add_bookmark(
                    h,
                    toc_bookmark_name(
                        "p", category_index, page_index
                    )
                )

                rendered_pages.add(
                    page_name
                )

            rendered_tabs = set()

            for page in items:

                tab = (
                    page.get("tab") or ""
                ).strip()

                if (
                    tab
                    and tab not in rendered_tabs
                ):
                    tab_index = len(rendered_tabs)
                    h = document.add_heading(
                        tab,
                        level=3
                    )

                    add_bookmark(
                        h,
                        toc_bookmark_name(
                            "t",
                            category_index,
                            page_index,
                            tab_index
                        )
                    )

                    rendered_tabs.add(
                        tab
                    )

                try:

                    section = (
                        generate_manual_section(
                            page
                        )
                    )
                    add_action_section(
                        document,
                        page.get(
                            "actions",
                            []
                        ),
                        section.get(
                            "button_descriptions",
                            []
                        ),
                        language
                    )
                    print(
                        page.get(
                            "actions",
                            []
                        )
                    )
                    print(
                        section.get(
                            "button_descriptions",
                            []
                        )
                    )
                    for heading_key, key in [
                        ("overview", "overview"),
                        ("business_value", "business_value"),
                        ("page_sections", "page_sections"),
                        ("fields", "field_descriptions"),
                        ("best_practices", "best_practices"),
                        ("restrictions", "restrictions")
                    ]:

                        data = normalize_ai_content(
                            section.get(
                                key,
                                []
                            )
                        )

                        if not data:
                            continue

                        document.add_heading(
                            text_for(language, heading_key),
                            level=5
                        )

                        if isinstance(
                            data,
                            list
                        ):
                            add_bullet_list(
                                document,
                                data
                            )
                        else:
                            document.add_paragraph(
                                str(data)
                            )

                except Exception as e:

                    print(
                        f"AI內容產生失敗: {e}"
                    )

    back_section = document.add_section(
        WD_SECTION.NEW_PAGE
    )
    set_blank_header_footer(back_section)
    back_section.top_margin = Cm(1.2)
    back_section.bottom_margin = Cm(1.2)
    add_back_cover(document, config, language)

    enforce_white_background_black_text(document)

    document.save(
        output_path
    )

    materialize_toc_page_numbers(
        output_path,
        grouped,
    )

    print(
        f"✅ 已產生：{output_path}"
    )
